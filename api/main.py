from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from sqlmodel import SQLModel, Field, Session, create_engine, select
from pydantic import BaseModel
import stripe
from fastapi import Request, Header, Depends
import jwt
from passlib.hash import bcrypt
import os, uuid, boto3
from typing import Optional, List


# ---------------- Structured Asset Generation (heuristic-based) ----------------

STOPWORDS = set('''a an and are as at be but by for from has have if in into is it its of on or such that the their there these this to was were will with you your i we they our not no do does did can should would could may might just about over under up down out very really more most less least than then so because while when where which who whom whose what how'''.split())

def split_sentences(text: str):
    # Basic sentence splitter
    text = text.replace("\n", " ").strip()
    parts = re.split(r'(?<=[.!?])\s+', text)
    return [p.strip() for p in parts if p.strip()]

def keywords(text: str, top_n=12):
    words = re.findall(r"[A-Za-z][A-Za-z'-]+", text.lower())
    words = [w for w in words if w not in STOPWORDS and len(w) > 2]
    c = Counter(words)
    return [w for w, _ in c.most_common(top_n)]

def title_options(text: str, asset_type: str):
    keys = keywords(text, 8)
    base = keys[:3] if keys else ["Insights", "Guide", "Basics"]
    patterns = [
        f"{asset_type.title()}: The {base[0].title()} Playbook",
        f"{base[0].title()} to {base[1].title()} — A {asset_type.title()} for Busy People",
        f"{asset_type.title()} Quickstart: {base[0].title()}, {base[1].title()}, and {base[2].title() if len(base)>2 else 'Essentials'}",
        f"{base[0].title()} Made Simple ({asset_type.title()})",
        f"The No-Fluff {asset_type.title()} on {base[0].title()}"
    ]
    # Deduplicate and ensure 5 items
    seen, out = set(), []
    for p in patterns:
        if p not in seen:
            seen.add(p); out.append(p)
    while len(out) < 5:
        out.append(f"{asset_type.title()} Essentials")
    return out[:5]

def cta_block(cta_type: str, cta_url: str):
    if not cta_type or not cta_url:
        return ""
    label = ""
    cta_type_lower = (cta_type or "").lower()
    if "webinar" in cta_type_lower and "replay" in cta_type_lower:
        label = "Watch the Webinar Replay"
    elif "webinar" in cta_type_lower:
        label = "Register for the Webinar"
    elif "schedule" in cta_type_lower or "call" in cta_type_lower:
        label = "Schedule a Call"
    elif "offer" in cta_type_lower or "checkout" in cta_type_lower or "buy" in cta_type_lower:
        label = "Claim the Offer"
    else:
        label = cta_type.title()
    return f"<div class='cta'><strong>Next Step:</strong> <a href='{cta_url}' target='_blank' rel='noreferrer'>{label}</a></div>"

DISCLAIMER_HTML = (
    "<p><em>This material is for informational and educational purposes only and is provided “as is.” "
    "It is not medical, legal, financial, or professional advice. Always consult a qualified professional for guidance specific to your situation. "
    "The publisher and author disclaim any liability for actions taken based on this content. All trademarks remain the property of their respective owners; "
    "references are for identification only and do not imply endorsement.</em></p>"
)

def cover_placeholder(title: str, subtitle: str = "Add your subtitle here"):
    return f\"\"\"<div class='cover' style='border:1px dashed #bbb; padding:16px; margin:12px 0'>
    <h1>{title}</h1>
    <p>{subtitle}</p>
    <p><small>[Upload your branded cover image here]</small></p>
</div>\"\"\"

def make_ebook_html(text: str, cta_type: str, cta_url: str):
    sents = split_sentences(text)
    keys = keywords(text, 10)
    ttl = keys[0].title() + " Guide" if keys else "Practical Guide"
    titles = title_options(text, "eBook")

    # Chunk sentences into 5–7 "chapters"
    n = max(5, min(7, max(5, len(sents)//6 or 5)))
    chunks = [sents[i::n] for i in range(n)]  # round-robin to balance
    chapters = []
    for idx, ch in enumerate(chunks, start=1):
        ch_text = " ".join(ch).strip()
        ck = keywords(ch_text, 1)
        ch_title = f"Chapter {idx}: {ck[0].title() if ck else f'Key Idea {idx}'}"
        chapters.append((ch_title, ch_text))

    html = [cover_placeholder(ttl)]
    html.append("<h2>Title Options</h2><ol>" + "".join([f"<li>{t}</li>" for t in titles]) + "</ol>")
    html.append("<h2>Introduction</h2><p>" + (" ".join(sents[:5]) or "This guide distills the core ideas into simple steps.") + "</p>")
    for ch_title, ch_text in chapters:
        html.append(f"<h2>{ch_title}</h2><p>{ch_text}</p>")
    # simple FAQ from top keywords
    html.append("<h2>FAQs</h2><ul>")
    for k in keys[:4]:
        html.append(f"<li><strong>What about {k}?</strong> <br/>Here's the short answer based on the video: focus on clarity, keep it simple, and apply the steps consistently.</li>")
    html.append("</ul>")
    html.append("<h2>Disclaimer</h2>" + DISCLAIMER_HTML)
    html.append(cta_block(cta_type, cta_url))
    return "\n".join(html), titles

def make_checklist_html(text: str, cta_type: str, cta_url: str):
    sents = split_sentences(text)
    candidates = [s for s in sents if len(s) < 180 and re.search(r'\\b(do|use|start|choose|add|remove|avoid|keep|set|try|plan|check|review|measure|track|calculate|define|decide|create|prepare|ensure|remember|consider|follow)\\b', s.lower())]
    if len(candidates) < 8:
        # fallback: pick short sentences
        candidates = [s for s in sents if len(s) < 140][:12]
    titles = title_options(text, "Checklist")
    html = [cover_placeholder("Action Checklist")]
    html.append("<h2>Title Options</h2><ol>" + "".join([f"<li>{t}</li>" for t in titles]) + "</ol>")
    html.append("<h2>Checklist</h2><ul>")
    for s in candidates[:18]:
        html.append(f"<li>☐ {s}</li>")
    html.append("</ul>")
    html.append("<h2>Disclaimer</h2>" + DISCLAIMER_HTML)
    html.append(cta_block(cta_type, cta_url))
    return "\n".join(html), titles

def make_cheatsheet_html(text: str, cta_type: str, cta_url: str):
    ks = keywords(text, 10)
    sents = split_sentences(text)
    short = [s for s in sents if len(s) < 120][:10]
    titles = title_options(text, "Cheat Sheet")
    html = [cover_placeholder("Cheat Sheet")]
    html.append("<h2>Title Options</h2><ol>" + "".join([f"<li>{t}</li>" for t in titles]) + "</ol>")
    html.append("<h2>Key Terms</h2><p>" + ", ".join([k.title() for k in ks]) + "</p>")
    html.append("<h2>Do</h2><ul>")
    for s in short[:5]:
        html.append(f"<li>{s}</li>")
    html.append("</ul><h2>Don't</h2><ul>")
    for s in short[5:10]:
        html.append(f"<li>{'Avoid: ' + s if not s.lower().startswith('avoid') else s}</li>")
    html.append("</ul>")
    html.append("<h2>Quick Tips</h2><ul>")
    for tip in ks[:5]:
        html.append(f"<li>Keep {tip} simple and consistent.</li>")
    html.append("</ul>")
    html.append("<h2>Disclaimer</h2>" + DISCLAIMER_HTML)
    html.append(cta_block(cta_type, cta_url))
    return "\n".join(html), titles

def make_onepager_html(text: str, cta_type: str, cta_url: str):
    sents = split_sentences(text)
    hook = sents[0] if sents else "Here’s the distilled, highly practical summary."
    ks = keywords(text, 8)
    bullets = sents[1:8] if len(sents) > 1 else []
    if not bullets:
        bullets = [f"Focus on {k} and apply it consistently." for k in ks[:6]]
    titles = title_options(text, "One-Page Summary")
    html = [cover_placeholder("One-Page Summary")]
    html.append("<h2>Title Options</h2><ol>" + "".join([f"<li>{t}</li>" for t in titles]) + "</ol>")
    html.append(f"<p><strong>Hook:</strong> {hook}</p>")
    html.append("<ul>")
    for b in bullets[:7]:
        html.append(f"<li>{b}</li>")
    html.append("</ul>")
    html.append("<h2>Disclaimer</h2>" + DISCLAIMER_HTML)
    html.append(cta_block(cta_type, cta_url))
    return "\n".join(html), titles


DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./app.db")
engine = create_engine(DATABASE_URL)

app = FastAPI(title="Video to Lead Magnets API")

origins = [os.getenv("ORIGINS", "http://localhost:3000")]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class Job(SQLModel, table=True):
    tenant_id: str = Field(index=True)
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    source_type: str
    source_ref: str
    status: str = "queued"

class Transcript(SQLModel, table=True):
    tenant_id: str = Field(index=True)
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    job_id: str = Field(index=True, foreign_key="job.id")
    raw_text: str = ""
    cleaned_text: str = ""

class Asset(SQLModel, table=True):
    tenant_id: str = Field(index=True)
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    job_id: str = Field(index=True, foreign_key="job.id")
    type: str
    title: str
    html: str = ""
    cta_type: Optional[str] = None
    cta_url: Optional[str] = None


def get_tenant_id(request: Request) -> str:
    tid = request.headers.get("X-Tenant-Id")
    if tid:
        return tid
    # Host header mapping
    host = (request.headers.get("host") or "").split(":")[0].lower()
    if host and host not in ("localhost", "127.0.0.1"):
        with Session(engine) as s:
            d = s.exec(select(TenantDomain).where(TenantDomain.hostname==host, TenantDomain.verified==True)).first()
            if d:
                return d.tenant_id
    return "demo-tenant"



def init_db():
    SQLModel.metadata.create_all(engine)

@app.on_event("startup")
def on_startup():
    init_db()

class CreateJobByUrl(BaseModel):
    video_url: str

class GenerateAssetsBody(BaseModel):
    job_id: str
    asset_types: List[str]
    cta_type: Optional[str] = None
    cta_url: Optional[str] = None

@app.get("/health")
def health():
    return {"ok": True}

@app.post("/jobs/url")
def create_job_by_url(body: CreateJobByUrl, request: Request):
    tenant_id = get_tenant_id(request)
    with Session(engine) as s:
        job = Job(tenant_id=tenant_id, source_type="url", source_ref=body.video_url, status="queued")
        s.add(job); s.commit(); s.refresh(job)
        return {"job_id": job.id, "status": job.status}

@app.post("/jobs/upload")
def create_job_by_upload(file: UploadFile = File(...)):
    s3 = boto3.client(
        "s3",
        endpoint_url=os.getenv("S3_ENDPOINT_URL"),
        aws_access_key_id=os.getenv("S3_ACCESS_KEY"),
        aws_secret_access_key=os.getenv("S3_SECRET_KEY"),
    )
    bucket = os.getenv("S3_BUCKET", "uploads")
    # Ensure bucket exists
    try:
        s3.head_bucket(Bucket=bucket)
    except Exception:
        s3.create_bucket(Bucket=bucket)
    key = f"uploads/{uuid.uuid4()}-{file.filename}"
    s3.upload_fileobj(file.file, bucket, key)
    with Session(engine) as s:
        job = Job(source_type="upload", source_ref=key, status="queued")
        s.add(job); s.commit(); s.refresh(job)
        return {"job_id": job.id, "status": job.status}

@app.get("/jobs/{job_id}")
def get_job(job_id: str):
    with Session(engine) as s:
        job = s.get(Job, job_id)
        if not job:
            raise HTTPException(404, "job not found")
        q = select(Transcript).where(Transcript.job_id == job_id)
        transcript = s.exec(q).first()
                if not job or job.tenant_id != tenant_id:
            raise HTTPException(404, "job not found")
        return {"job": job, "transcript": transcript}

@app.post("/generate")
def generate_assets(body: GenerateAssetsBody):
    with Session(engine) as s:
        job = s.get(Job, body.job_id)
        if not job:
            raise HTTPException(404, "job not found")
        job.status = "drafting"
        s.add(job); s.commit()

        # fetch transcript
        q = select(Transcript).where(Transcript.job_id == job.id)
        tr = s.exec(q).first()
        source_text = (tr.cleaned_text or tr.raw_text) if tr else ""

        created_ids = []
        for t in body.asset_types:
            t_lower = t.lower().strip()
            if t_lower in ["ebook", "e-book", "book"]:
                html, titles = make_ebook_html(source_text, body.cta_type, body.cta_url)
                title = titles[0]
            elif t_lower in ["checklist"]:
                html, titles = make_checklist_html(source_text, body.cta_type, body.cta_url)
                title = titles[0]
            elif t_lower in ["cheat sheet", "cheatsheet", "cheat-sheet"]:
                html, titles = make_cheatsheet_html(source_text, body.cta_type, body.cta_url)
                title = titles[0]
            elif t_lower in ["one-page summary", "one pager", "one-pager", "summary"]:
                html, titles = make_onepager_html(source_text, body.cta_type, body.cta_url)
                title = titles[0]
            else:
                # default to one-pager
                html, titles = make_onepager_html(source_text, body.cta_type, body.cta_url)
                title = titles[0]

            asset = Asset(tenant_id=tenant_id, job_id=job.id, type=t_lower, title=title, html=html, cta_type=body.cta_type, cta_url=body.cta_url)
            s.add(asset); s.commit(); s.refresh(asset)
            created_ids.append(asset.id)

        job.status = "ready"
        s.add(job); s.commit()
        log_action(tenant_id, "generate_assets", f"{len(body.asset_types)} assets", user.get("uid"));
        return {"job_id": job.id, "asset_ids": created_ids}

from docx import Document
from fastapi.responses import StreamingResponse
import tempfile, re

@app.get("/assets/{asset_id}")
def get_asset(asset_id: str):
    with Session(engine) as s:
        a = s.get(Asset, asset_id)
        if not a or a.tenant_id != tenant_id:
            raise HTTPException(404, "asset not found")
        return a

@app.get("/export/docx/{asset_id}")
def export_docx(asset_id: str):
    with Session(engine) as s:
        a = s.get(Asset, asset_id)
        if not a or a.tenant_id != tenant_id:
            raise HTTPException(404, "asset not found")
        doc = Document()
        doc.add_heading(a.title, 0)
        text = re.sub("<[^<]+?>", "", a.html)
        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        disclaimer = ("This material is for informational and educational purposes only and is provided “as is.” "
                      "It is not professional advice. Consult a qualified professional for guidance specific to your situation. "
                      "The publisher and author disclaim any liability for actions taken based on this content.")
        doc.add_page_break()
        doc.add_heading("Disclaimer", level=1)
        doc.add_paragraph(disclaimer)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            return StreamingResponse(open(tmp.name, "rb"),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={a.title}.docx"})


from fastapi.responses import StreamingResponse
from weasyprint import HTML, CSS
import io

@app.get("/export/pdf/{asset_id}")
def export_pdf(asset_id: str):
    with Session(engine) as s:
        a = s.get(Asset, asset_id)
        if not a or a.tenant_id != tenant_id:
            raise HTTPException(404, "asset not found")
        # Minimal CSS for nicer output + cover separation
        base_css = CSS(string=\"\"\"
            @page { size: A4; margin: 24mm 18mm; }
            h1, h2, h3 { page-break-after: avoid; }
            .cover { border: 2px dashed #bbb; padding: 20px; margin: 12px 0; text-align:center; }
            .cta { margin-top: 16px; padding: 12px; border: 1px solid #ddd; }
            ul { margin: 0 0 0 18px; }
            li { margin: 6px 0; }
        \"\"\")
        html = f\"\"\"<html><head><meta charset="utf-8"></head><body>
        {a.html}
        </body></html>\"\"\"
        pdf_bytes = HTML(string=html).write_pdf(stylesheets=[base_css])
        return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf",
                                 headers={"Content-Disposition": f"attachment; filename={a.title}.pdf"})


class UpdateAssetBody(BaseModel):
    title: str
    html: str
    cta_type: Optional[str] = None
    cta_url: Optional[str] = None

@app.put("/assets/{asset_id}")
def update_asset(asset_id: str, body: UpdateAssetBody):
    with Session(engine) as s:
        a = s.get(Asset, asset_id)
        if not a or a.tenant_id != tenant_id:
            raise HTTPException(404, "asset not found")
        a.title = body.title
        a.html = body.html
        a.cta_type = body.cta_type
        a.cta_url = body.cta_url
        s.add(a); s.commit(); s.refresh(a)
        log_action(tenant_id, "update_asset", f"asset_id={asset_id}", user.get("uid"))
        return a


class BillingCustomer(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    tenant_id: str = Field(index=True)
    email: str
    stripe_customer_id: str | None = None
    stripe_subscription_id: str | None = None
    subscription_status: str | None = None


# ---------------- Stripe Billing ----------------
class StripeCheckoutBody(BaseModel):
    tenant_id: str
    email: str
    price_id: str | None = None

@app.post("/billing/stripe/checkout")
def stripe_checkout(body: StripeCheckoutBody, request: Request):
    secret = os.getenv("STRIPE_SECRET_KEY")
    price_id = body.price_id or os.getenv("STRIPE_PRICE_ID")
    public_app = os.getenv("PUBLIC_APP_URL", "http://localhost:3000")
    if not secret or not price_id:
        raise HTTPException(400, "Stripe not configured")
    user = current_user(request)
    require_role(user, "admin", "owner")
    user = current_user(request)
    require_role(user, "admin", "owner")
    stripe.api_key = secret

    # Ensure a BillingCustomer row
    with Session(engine) as s:
        q = select(BillingCustomer).where(BillingCustomer.tenant_id==body.tenant_id, BillingCustomer.email==body.email)
        bc = s.exec(q).first()
        if not bc:
            bc = BillingCustomer(tenant_id=body.tenant_id, email=body.email)
            s.add(bc); s.commit(); s.refresh(bc)

    session = stripe.checkout.Session.create(
        mode="subscription",
        line_items=[{"price": price_id, "quantity": 1}],
        customer_email=body.email,
        success_url=f"{public_app}/?checkout=success",
        cancel_url=f"{public_app}/?checkout=cancel",
        metadata={"tenant_id": body.tenant_id, "billing_customer_id": bc.id},
    )
    return {"url": session.url}

from fastapi import Header
@app.post("/billing/stripe/webhook")
def stripe_webhook(request: Request):
    payload = request.body()
    sig_header = request.headers.get("stripe-signature")
    webhook_secret = os.getenv("STRIPE_WEBHOOK_SECRET")
    if not webhook_secret:
        raise HTTPException(400, "Webhook not configured")
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except Exception as e:
        raise HTTPException(400, f"Webhook error: {e}")

    et = event["type"]
    data = event["data"]["object"]

    with Session(engine) as s:
        if et == "checkout.session.completed":
            meta = data.get("metadata", {}) or {}
            bc_id = meta.get("billing_customer_id")
            if bc_id:
                bc = s.get(BillingCustomer, bc_id)
                if bc:
                    bc.stripe_customer_id = data.get("customer")
                    s.add(bc); s.commit()
        elif et == "checkout.session.completed" and (data.get("mode") == "payment"):
            tenant_id = (data.get("metadata") or {}).get("tenant_id")
            with Session(engine) as s:
                uc = s.exec(select(UsageCredits).where(UsageCredits.tenant_id==tenant_id)).first()
                if not uc:
                    uc = UsageCredits(tenant_id=tenant_id, balance=0)
                pack = int(os.getenv("TOPUP_PACK_CREDITS", "50"))
                uc.balance += pack
                s.add(uc); s.commit()
        elif et == "customer.subscription.created" or et == "customer.subscription.updated":
            tenant_id = (data.get("metadata") or {}).get("tenant_id")
            sub_id = data.get("id")
            status = data.get("status")
            cust_id = data.get("customer")
            # Update all BC rows for this tenant/customer
            q = select(BillingCustomer).where(BillingCustomer.stripe_customer_id==cust_id)
            for bc in s.exec(q).all():
                bc.stripe_subscription_id = sub_id
                bc.subscription_status = status
                s.add(bc)
            s.commit()

    return {"received": True}


@app.post("/billing/stripe/portal")
def stripe_portal(body: StripeCheckoutBody, request: Request):
    secret = os.getenv("STRIPE_SECRET_KEY")
    public_app = os.getenv("PUBLIC_APP_URL", "http://localhost:3000")
    if not secret:
        raise HTTPException(400, "Stripe not configured")
    user = current_user(request)
    require_role(user, "admin", "owner")
    stripe.api_key = secret
    with Session(engine) as s:
        q = select(BillingCustomer).where(BillingCustomer.tenant_id==body.tenant_id, BillingCustomer.email==body.email)
        bc = s.exec(q).first()
        if not bc or not bc.stripe_customer_id:
            raise HTTPException(404, "Customer not found. Complete checkout first.")
    portal = stripe.billing_portal.Session.create(
        customer=bc.stripe_customer_id,
        return_url=public_app
    )
    return {"url": portal.url}


class TenantBranding(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    tenant_id: str = Field(index=True, unique=True)
    name: str = "Your Brand"
    primary_color: str = "#0ea5e9"
    accent_color: str = "#22c55e"
    logo_url: str | None = None
    domain: str | None = None


class BrandingBody(BaseModel):
    name: str | None = None
    primary_color: str | None = None
    accent_color: str | None = None
    logo_url: str | None = None
    domain: str | None = None

@app.get("/tenant/branding")
def get_branding(request: Request):
    tenant_id = get_tenant_id(request)
    with Session(engine) as s:
        q = select(TenantBranding).where(TenantBranding.tenant_id==tenant_id)
        b = s.exec(q).first()
        if not b:
            b = TenantBranding(tenant_id=tenant_id)
            s.add(b); s.commit(); s.refresh(b)
        return b

@app.put("/tenant/branding")
def put_branding(body: BrandingBody, request: Request):
    tenant_id = get_tenant_id(request)
    with Session(engine) as s:
        q = select(TenantBranding).where(TenantBranding.tenant_id==tenant_id)
        b = s.exec(q).first()
        if not b:
            b = TenantBranding(tenant_id=tenant_id)
        if body.name is not None: b.name = body.name
        if body.primary_color is not None: b.primary_color = body.primary_color
        if body.accent_color is not None: b.accent_color = body.accent_color
        if body.logo_url is not None: b.logo_url = body.logo_url
        if body.domain is not None: b.domain = body.domain
        s.add(b); s.commit(); s.refresh(b)
        return b


class User(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    email: str = Field(unique=True, index=True)
    password_hash: str
    email_verified: bool = False

class Membership(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    user_id: str = Field(index=True, foreign_key="user.id")
    tenant_id: str = Field(index=True)
    role: str = "editor"  # owner|admin|editor

class UsageCredits(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    tenant_id: str = Field(index=True, unique=True)
    balance: int = 100  # starter credits

class AuditLog(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    tenant_id: str = Field(index=True)
    user_id: str | None = None
    action: str
    details: str = ""


JWT_SECRET = os.getenv("JWT_SECRET", "devsecret")
CREDITS_PER_ASSET = int(os.getenv("CREDITS_PER_ASSET", "1"))

def log_action(tenant_id: str, action: str, details: str = "", user_id: str | None = None):
    with Session(engine) as s:
        s.add(AuditLog(tenant_id=tenant_id, user_id=user_id, action=action, details=details))
        s.commit()

def make_token(user_id: str, email: str, tenant_id: str, role: str):
    payload = {"uid": user_id, "email": email, "tenant_id": tenant_id, "role": role}
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

def current_user(request: Request):
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(401, "Missing token")
    token = auth.split(" ", 1)[1]
    try:
        data = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
    except Exception:
        raise HTTPException(401, "Invalid token")
    return data  # dict with uid, email, tenant_id, role

def require_role(user, *allowed):
    role = (user.get("role") or "").lower()
    order = {"owner":3, "admin":2, "editor":1}
    # allow if role is in allowed set or higher privilege
    allowed_set = set([r.lower() for r in allowed])
    if role in allowed_set:
        return True
    # owner is highest; admin above editor
    if "owner" in allowed_set and role == "owner": return True
    if "admin" in allowed_set and role in ["owner", "admin"]: return True
    if "editor" in allowed_set and role in ["owner", "admin", "editor"]: return True
    raise HTTPException(403, "Insufficient role")


class SignupBody(BaseModel):
    email: str
    password: str
    tenant_id: str
    role: str = "owner"

class LoginBody(BaseModel):
    email: str
    password: str
    tenant_id: str

@app.post("/auth/signup")
def signup(body: SignupBody):
    with Session(engine) as s:
        existing = s.exec(select(User).where(User.email == body.email)).first()
        if existing:
            raise HTTPException(400, "Email already exists")
        user = User(email=body.email, password_hash=bcrypt.hash(body.password))
        s.add(user); s.commit(); s.refresh(user)
        # create membership
        mem = Membership(user_id=user.id, tenant_id=body.tenant_id, role=body.role.lower())
        s.add(mem)
        # give credits row if none
        if not s.exec(select(UsageCredits).where(UsageCredits.tenant_id==body.tenant_id)).first():
            s.add(UsageCredits(tenant_id=body.tenant_id, balance=100))
        s.commit()
        token = make_token(user.id, user.email, body.tenant_id, body.role.lower())
        log_action(body.tenant_id, "signup", f"user {user.email} role {body.role}", user.id)
        return {"token": token, "role": body.role.lower(), "tenant_id": body.tenant_id}

@app.post("/auth/login")
def login(body: LoginBody):
    with Session(engine) as s:
        user = s.exec(select(User).where(User.email == body.email)).first()
        if not user or not bcrypt.verify(body.password, user.password_hash):
            raise HTTPException(401, "Invalid credentials")
        mem = s.exec(select(Membership).where(Membership.user_id==user.id, Membership.tenant_id==body.tenant_id)).first()
        if not mem:
            raise HTTPException(403, "No access to tenant")
        token = make_token(user.id, user.email, mem.tenant_id, mem.role)
        log_action(mem.tenant_id, "login", f"user {user.email}", user.id)
        return {"token": token, "role": mem.role, "tenant_id": mem.tenant_id, "email_verified": u.email_verified}


@app.get("/usage/balance")
def usage_balance(request: Request):
    tenant_id = get_tenant_id(request)
    with Session(engine) as s:
        uc = s.exec(select(UsageCredits).where(UsageCredits.tenant_id==tenant_id)).first()
        if not uc:
            uc = UsageCredits(tenant_id=tenant_id, balance=0); s.add(uc); s.commit(); s.refresh(uc)
        return {"tenant_id": tenant_id, "balance": uc.balance}

class TopUpBody(BaseModel):
    amount: int

@app.post("/usage/topup")
def usage_topup(body: TopUpBody, request: Request):
    tenant_id = get_tenant_id(request)
    user = current_user(request)
    require_role(user, "admin", "owner")
    with Session(engine) as s:
        uc = s.exec(select(UsageCredits).where(UsageCredits.tenant_id==tenant_id)).first()
        if not uc:
            uc = UsageCredits(tenant_id=tenant_id, balance=0)
        uc.balance += max(0, int(body.amount))
        s.add(uc); s.commit(); s.refresh(uc)
        log_action(tenant_id, "credits_topup", f"+{body.amount}", user.get("uid"))
        return {"tenant_id": tenant_id, "balance": uc.balance}

@app.get("/audit")
def audit(request: Request, limit: int = 100):
    tenant_id = get_tenant_id(request)
    user = current_user(request)
    require_role(user, "admin", "owner")
    with Session(engine) as s:
        rows = s.exec(select(AuditLog).where(AuditLog.tenant_id==tenant_id)).all()
        rows = rows[-limit:]
        return [{"ts": r.id, "action": r.action, "details": r.details, "user_id": r.user_id} for r in rows]

class EmailToken(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    tenant_id: str = Field(index=True)
    user_id: str = Field(index=True)
    email: str
    token: str = Field(index=True, unique=True)
    kind: str  # 'verify' or 'reset'

def send_postmark(to_email: str, subject: str, html: str):
    import requests
    token = os.getenv("POSTMARK_SERVER_TOKEN")
    sender = os.getenv("POSTMARK_FROM", "hello@example.com")
    if not token:
        # In dev, just print
        print("[DEV EMAIL]", to_email, subject, html[:140])
        return True
    r = requests.post(
        "https://api.postmarkapp.com/email",
        headers={"X-Postmark-Server-Token": token, "Accept":"application/json"},
        json={"From": sender, "To": to_email, "Subject": subject, "HtmlBody": html}
    )
    return r.status_code in (200, 201)

@app.post("/auth/send-verify")
def send_verify(request: Request):
    user = current_user(request)
    tenant_id = user["tenant_id"]
    email = user["email"]
    with Session(engine) as s:
        u = s.exec(select(User).where(User.email==email)).first()
        if not u:
            raise HTTPException(404, "User not found")
        token = uuid.uuid4().hex
        s.add(EmailToken(tenant_id=tenant_id, user_id=u.id, email=u.email, token=token, kind="verify"))
        s.commit()
        app_url = os.getenv("PUBLIC_APP_URL", "http://localhost:3000")
        link = f"{app_url}/verify?token={token}"
        send_postmark(u.email, "Verify your email", f"<p>Click to verify:</p><p><a href='{link}'>{link}</a></p>")
        log_action(tenant_id, "email_send_verify", u.email, u.id)
        return {"ok": True}

@app.post("/auth/verify")
def verify_email(token: str):
    with Session(engine) as s:
        row = s.exec(select(EmailToken).where(EmailToken.token==token, EmailToken.kind=='verify')).first()
        if not row:
            raise HTTPException(400, "Invalid token")
        u = s.get(User, row.user_id)
        if not u:
            raise HTTPException(404, "User missing")
        u.email_verified = True
        s.add(u); s.commit()
        log_action(row.tenant_id, "email_verified", u.email, u.id)
        return {"ok": True}

class ForgotBody(BaseModel):
    email: str
    tenant_id: str

@app.post("/auth/forgot")
def forgot(body: ForgotBody):
    with Session(engine) as s:
        u = s.exec(select(User).where(User.email==body.email)).first()
        if not u:
            return {"ok": True}  # don't leak existence
        token = uuid.uuid4().hex
        s.add(EmailToken(tenant_id=body.tenant_id, user_id=u.id, email=u.email, token=token, kind="reset"))
        s.commit()
        app_url = os.getenv("PUBLIC_APP_URL", "http://localhost:3000")
        link = f"{app_url}/reset?token={token}"
        send_postmark(u.email, "Reset your password", f"<p>Reset link:</p><p><a href='{link}'>{link}</a></p>")
        log_action(body.tenant_id, "email_send_reset", u.email, u.id)
        return {"ok": True}

class ResetBody(BaseModel):
    token: str
    new_password: str

@app.post("/auth/reset")
def reset(body: ResetBody):
    with Session(engine) as s:
        row = s.exec(select(EmailToken).where(EmailToken.token==body.token, EmailToken.kind=='reset')).first()
        if not row:
            raise HTTPException(400, "Invalid token")
        u = s.get(User, row.user_id)
        if not u:
            raise HTTPException(404, "User missing")
        u.password_hash = bcrypt.hash(body.new_password)
        s.add(u); s.commit()
        log_action(row.tenant_id, "password_reset", u.email, u.id)
        return {"ok": True}

class TenantDomain(SQLModel, table=True):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()), primary_key=True)
    tenant_id: str = Field(index=True)
    hostname: str = Field(unique=True, index=True)
    verify_token: str = Field(default_factory=lambda: uuid.uuid4().hex)
    verified: bool = False

class DomainBody(BaseModel):
    hostname: str

@app.post("/domains/add")
def add_domain(body: DomainBody, request: Request):
    user = current_user(request); require_role(user, "admin", "owner")
    tenant_id = get_tenant_id(request)
    with Session(engine) as s:
        existing = s.exec(select(TenantDomain).where(TenantDomain.hostname==body.hostname)).first()
        if existing:
            raise HTTPException(400, "Hostname already in use")
        d = TenantDomain(tenant_id=tenant_id, hostname=body.hostname)
        s.add(d); s.commit(); s.refresh(d)
        # In DNS: add TXT record verify_token at _vtlm.{hostname}
        return {"hostname": d.hostname, "txt_record": f"_vtlm.{d.hostname} = {d.verify_token}"}

@app.post("/domains/verify")
def verify_domain(body: DomainBody, request: Request):
    user = current_user(request); require_role(user, "admin", "owner")
    tenant_id = get_tenant_id(request)
    with Session(engine) as s:
        d = s.exec(select(TenantDomain).where(TenantDomain.hostname==body.hostname, TenantDomain.tenant_id==tenant_id)).first()
        if not d:
            raise HTTPException(404, "Domain not found")
        # In a real system, we'd check DNS TXT. Here, assume verified when endpoint is called.
        d.verified = True
        s.add(d); s.commit(); s.refresh(d)
        return {"ok": True, "hostname": d.hostname, "verified": d.verified}


@app.get("/me")
def me(request: Request):
    user_data = current_user(request)
    with Session(engine) as s:
        u = s.get(User, user_data["uid"])
        if not u:
            raise HTTPException(404, "User not found")
        return {"email": u.email, "email_verified": u.email_verified, "tenant_id": user_data["tenant_id"], "role": user_data["role"]}
